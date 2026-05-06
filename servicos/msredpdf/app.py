from __future__ import annotations

import asyncio
import io
import json
import os
import re
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, AsyncIterator

import fitz
import httpx
from docx import Document
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse


APP_DIR = Path(__file__).resolve().parent
STATIC_DIR = APP_DIR / "static"


def env_int(name: str, default: int) -> int:
    try:
        return int(os.getenv(name, str(default)))
    except ValueError:
        return default


def env_bool(name: str, default: bool) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "on", "sim"}


DATA_DIR = Path(os.getenv("MSREDPDF_DATA_DIR", "/var/lib/msredpdf"))
UPLOAD_DIR = DATA_DIR / "uploads"
RESULT_DIR = DATA_DIR / "results"
PROXY_BASE_URL = os.getenv("MSREDPDF_PROXY_BASE_URL", "http://127.0.0.1:8080/v1").rstrip("/")
PROXY_API_KEY = os.getenv("MSREDPDF_PROXY_API_KEY", "").strip()
DEFAULT_MODEL = os.getenv("MSREDPDF_MODEL", "mistral-large-latest")
REVIEW_MODEL = os.getenv("MSREDPDF_REVIEW_MODEL", "glm-5.1")
REDACT_BEFORE_AI = env_bool("MSREDPDF_REDACT_BEFORE_AI", True)
MAX_UPLOAD_MB = env_int("MSREDPDF_MAX_UPLOAD_MB", 80)
MAX_ANALYSIS_PAGES = env_int("MSREDPDF_MAX_ANALYSIS_PAGES", 80)
MAX_CHARS_WHOLE = env_int("MSREDPDF_MAX_CHARS_WHOLE", 45000)
MAX_CHARS_PER_PAGE = env_int("MSREDPDF_MAX_CHARS_PER_PAGE", 12000)
REPORT_TOKENS = env_int("MSREDPDF_REPORT_TOKENS", 5200)
PAGE_NOTE_TOKENS = env_int("MSREDPDF_PAGE_NOTE_TOKENS", 1000)
REQUEST_TIMEOUT = env_int("MSREDPDF_REQUEST_TIMEOUT", 180)
ENABLE_OCR = env_bool("MSREDPDF_ENABLE_OCR", True)
OCR_LANG = os.getenv("MSREDPDF_OCR_LANG", "por+eng").strip() or "por+eng"
OCR_DPI = env_int("MSREDPDF_OCR_DPI", 200)
OCR_MIN_CHARS = env_int("MSREDPDF_OCR_MIN_CHARS", 40)
OCR_MAX_PAGES = env_int("MSREDPDF_OCR_MAX_PAGES", 40)
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

RECOMMENDED_MODELS = [
    {"id": "mistral-large-latest", "label": "Mistral Large", "role": "analise profunda"},
    {"id": "mistral-medium-latest", "label": "Mistral Medium", "role": "equilibrado"},
    {"id": "kimi-k2.6", "label": "Kimi K2.6", "role": "raciocinio longo"},
    {"id": "glm-5.1", "label": "GLM 5.1", "role": "segunda opiniao"},
    {"id": "qwen3.5:397b", "label": "Qwen 3.5 397B", "role": "analise robusta"},
    {"id": "deepseek-v3.2", "label": "DeepSeek V3.2", "role": "sintese tecnica"},
]


@dataclass
class Job:
    id: str
    filename: str
    model: str
    mode: str
    document_type: str
    status: str = "queued"
    created_at: float = field(default_factory=time.time)
    updated_at: float = field(default_factory=time.time)
    progress: int = 0
    events: list[dict[str, Any]] = field(default_factory=list)
    report_draft: str = ""
    result: dict[str, Any] | None = None
    error: str | None = None
    file_path: str | None = None


app = FastAPI(title="MS RED PDF", version="0.1.0")
jobs: dict[str, Job] = {}


def ensure_dirs() -> None:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    RESULT_DIR.mkdir(parents=True, exist_ok=True)


def public_job(job: Job) -> dict[str, Any]:
    return {
        "id": job.id,
        "filename": job.filename,
        "model": job.model,
        "mode": job.mode,
        "document_type": job.document_type,
        "status": job.status,
        "progress": job.progress,
        "report_draft": job.report_draft,
        "created_at": job.created_at,
        "updated_at": job.updated_at,
        "result": job.result,
        "error": job.error,
    }


def public_saved_job(job_id: str, result: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": job_id,
        "filename": result.get("filename", "documento"),
        "model": result.get("model", ""),
        "mode": result.get("mode", ""),
        "document_type": result.get("document", {}).get("type", ""),
        "status": "completed",
        "progress": 100,
        "created_at": result.get("created_at"),
        "updated_at": result.get("finished_at"),
        "result": result,
        "error": None,
    }


def add_event(job: Job, stage: str, message: str, progress: int | None = None, detail: Any = None) -> None:
    if progress is not None:
        job.progress = max(0, min(100, progress))
    job.updated_at = time.time()
    event = {
        "seq": len(job.events) + 1,
        "time": job.updated_at,
        "stage": stage,
        "status": job.status,
        "message": message,
        "progress": job.progress,
        "detail": detail,
    }
    job.events.append(event)


def redact_for_lgpd(text: str) -> str:
    if not REDACT_BEFORE_AI:
        return text
    patterns = [
        (r"\b\d{3}\.?\d{3}\.?\d{3}-?\d{2}\b", "[CPF_REMOVIDO]"),
        (r"\b\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}\b", "[CNPJ_REMOVIDO]"),
        (r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b", "[EMAIL_REMOVIDO]"),
        (r"(?<!\d)(?:\+?55\s*)?(?:\(?\d{2}\)?\s*)?\d{4,5}[-\s]?\d{4}(?!\d)", "[TELEFONE_REMOVIDO]"),
    ]
    redacted = text
    for pattern, replacement in patterns:
        redacted = re.sub(pattern, replacement, redacted)
    return redacted


def trim_text(text: str, limit: int) -> str:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[TRECHO CORTADO POR LIMITE DE CONTEXTO]"


def extract_pdf(path: Path, job: Job) -> dict[str, Any]:
    pages: list[dict[str, Any]] = []
    doc = fitz.open(path)
    metadata = dict(doc.metadata or {})
    total_pages = len(doc)
    add_event(job, "extract", f"PDF aberto com {total_pages} pagina(s).", 12, {"pages": total_pages})
    for index, page in enumerate(doc, start=1):
        text = page.get_text("text") or ""
        text = text.replace("\x00", "").strip()
        pages.append({"page": index, "chars": len(text), "text": text})
        progress = 12 + int((index / max(total_pages, 1)) * 23)
        add_event(job, "extract", f"Pagina {index}/{total_pages} extraida com {len(text)} caractere(s).", progress)
    doc.close()
    total_chars = sum(page["chars"] for page in pages)
    empty_pages = sum(1 for page in pages if page["chars"] < 40)
    return {
        "document_type": "pdf",
        "unit_label": "pagina",
        "metadata": metadata,
        "pages": pages,
        "total_pages": total_pages,
        "total_chars": total_chars,
        "empty_pages": empty_pages,
        "ocr_applied": False,
        "ocr_pages": [],
    }


def refresh_extraction_counts(extraction: dict[str, Any]) -> None:
    extraction["total_chars"] = sum(page["chars"] for page in extraction["pages"])
    extraction["empty_pages"] = sum(1 for page in extraction["pages"] if page["chars"] < OCR_MIN_CHARS)


def pdf_needs_ocr(extraction: dict[str, Any]) -> bool:
    if extraction["document_type"] != "pdf" or not ENABLE_OCR:
        return False
    if extraction["total_chars"] == 0:
        return True
    return any(page["chars"] < OCR_MIN_CHARS for page in extraction["pages"])


def apply_pdf_ocr(path: Path, job: Job, extraction: dict[str, Any]) -> dict[str, Any]:
    try:
        import pytesseract
        from PIL import Image
    except Exception as exc:
        raise RuntimeError(
            "OCR necessario, mas as dependencias pytesseract/Pillow nao estao instaladas no backend."
        ) from exc

    candidates = [page for page in extraction["pages"] if page["chars"] < OCR_MIN_CHARS]
    if not candidates:
        return extraction
    limited = candidates[:OCR_MAX_PAGES]
    skipped = len(candidates) - len(limited)
    add_event(
        job,
        "ocr",
        f"OCR ativado para {len(limited)} pagina(s) com pouco texto.",
        36,
        {"lang": OCR_LANG, "dpi": OCR_DPI, "skipped": skipped},
    )

    doc = fitz.open(path)
    matrix = fitz.Matrix(OCR_DPI / 72, OCR_DPI / 72)
    ocr_pages: list[int] = []
    try:
        total = max(len(limited), 1)
        for index, target in enumerate(limited, start=1):
            page_number = int(target["page"])
            page = doc[page_number - 1]
            progress = 36 + int((index - 1) / total * 9)
            add_event(job, "ocr", f"Lendo imagem da pagina {page_number} com OCR ({index}/{total}).", progress)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(image, lang=OCR_LANG, config="--psm 6") or ""
            text = text.replace("\x00", "").strip()
            if text:
                target["text"] = text
                target["chars"] = len(text)
                target["source"] = "ocr"
                ocr_pages.append(page_number)
                add_event(job, "ocr", f"Pagina {page_number} convertida por OCR com {len(text)} caractere(s).", progress + 1)
            else:
                add_event(job, "ocr", f"OCR nao encontrou texto util na pagina {page_number}.", progress + 1)
    finally:
        doc.close()

    extraction["ocr_applied"] = bool(ocr_pages)
    extraction["ocr_pages"] = ocr_pages
    extraction["ocr_lang"] = OCR_LANG
    extraction["ocr_dpi"] = OCR_DPI
    extraction["ocr_skipped_pages"] = skipped
    refresh_extraction_counts(extraction)
    if ocr_pages:
        add_event(job, "ocr", f"OCR finalizado em {len(ocr_pages)} pagina(s).", 45, {"pages": ocr_pages})
    return extraction


def extract_docx(path: Path, job: Job) -> dict[str, Any]:
    doc = Document(path)
    props = doc.core_properties
    metadata = {
        "author": props.author,
        "title": props.title,
        "subject": props.subject,
        "created": props.created.isoformat() if props.created else "",
        "modified": props.modified.isoformat() if props.modified else "",
    }
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    blocks: list[dict[str, Any]] = []
    current: list[str] = []
    current_len = 0
    for part in parts:
        if current and current_len + len(part) > 3500:
            text = "\n".join(current).strip()
            blocks.append({"page": len(blocks) + 1, "chars": len(text), "text": text})
            current = []
            current_len = 0
        current.append(part)
        current_len += len(part) + 1
    if current:
        text = "\n".join(current).strip()
        blocks.append({"page": len(blocks) + 1, "chars": len(text), "text": text})

    total_blocks = len(blocks)
    add_event(job, "extract", f"DOCX aberto com {total_blocks} bloco(s) de texto.", 12, {"blocks": total_blocks})
    for index, block in enumerate(blocks, start=1):
        progress = 12 + int((index / max(total_blocks, 1)) * 23)
        add_event(job, "extract", f"Bloco {index}/{total_blocks} extraido com {block['chars']} caractere(s).", progress)

    total_chars = sum(block["chars"] for block in blocks)
    empty_blocks = sum(1 for block in blocks if block["chars"] < 40)
    return {
        "document_type": "docx",
        "unit_label": "bloco",
        "metadata": metadata,
        "pages": blocks,
        "total_pages": total_blocks,
        "total_chars": total_chars,
        "empty_pages": empty_blocks,
        "ocr_applied": False,
        "ocr_pages": [],
    }


def extract_document(path: Path, job: Job) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path, job)
    if suffix == ".docx":
        return extract_docx(path, job)
    raise RuntimeError("Formato nao suportado. Envie PDF ou DOCX.")


def choose_strategy(mode: str, extraction: dict[str, Any]) -> str:
    if mode in {"whole", "pages"}:
        return mode
    if extraction["total_pages"] <= 6 and extraction["total_chars"] <= MAX_CHARS_WHOLE:
        return "whole"
    return "pages"


async def call_proxy(model: str, messages: list[dict[str, str]], max_tokens: int = 1600) -> str:
    headers = {"Content-Type": "application/json"}
    if PROXY_API_KEY:
        headers["Authorization"] = f"Bearer {PROXY_API_KEY}"
    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.15,
        "max_tokens": max_tokens,
    }
    async with httpx.AsyncClient(timeout=REQUEST_TIMEOUT) as client:
        response = await client.post(f"{PROXY_BASE_URL}/chat/completions", headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
    return data["choices"][0]["message"]["content"]


def stream_delta_from_chunk(data: dict[str, Any]) -> str:
    choices = data.get("choices") or []
    if not choices:
        return ""
    choice = choices[0] or {}
    delta = choice.get("delta") or {}
    if isinstance(delta, dict):
        content = delta.get("content")
        if isinstance(content, str):
            return content
    message = choice.get("message") or {}
    if isinstance(message, dict):
        content = message.get("content")
        if isinstance(content, str):
            return content
    return ""


def add_report_delta(job: Job, delta: str, progress: int) -> None:
    if not delta:
        return
    job.report_draft += delta
    add_event(job, "report-delta", "Trecho do relatorio recebido.", progress, {"delta": delta})


async def call_proxy_stream(
    model: str,
    messages: list[dict[str, str]],
    job: Job,
    max_tokens: int,
    progress_start: int,
    progress_end: int,
) -> str:
    headers = {"Content-Type": "application/json"}
    if PROXY_API_KEY:
        headers["Authorization"] = f"Bearer {PROXY_API_KEY}"
    payload = {
        "model": model,
        "messages": messages,
        "temperature": 0.15,
        "max_tokens": max_tokens,
        "stream": True,
    }
    chunks: list[str] = []
    pending: list[str] = []
    last_emit = time.monotonic()
    total_chars = 0
    prefix_buffer = ""
    stream_started = False

    async def flush(force: bool = False) -> None:
        nonlocal last_emit, total_chars
        if not pending:
            return
        now = time.monotonic()
        joined = "".join(pending)
        if not force and len(joined) < 24 and now - last_emit < 0.04:
            return
        pending.clear()
        for index in range(0, len(joined), 32):
            piece = joined[index : index + 32]
            total_chars += len(piece)
            estimated = min(progress_end, progress_start + int(total_chars / 450))
            add_report_delta(job, piece, estimated)
        last_emit = now

    try:
        async with httpx.AsyncClient(timeout=REQUEST_TIMEOUT) as client:
            async with client.stream("POST", f"{PROXY_BASE_URL}/chat/completions", headers=headers, json=payload) as response:
                response.raise_for_status()
                async for line in response.aiter_lines():
                    if not line:
                        continue
                    if line.startswith("data:"):
                        line = line[5:].strip()
                    if line == "[DONE]":
                        break
                    try:
                        data = json.loads(line)
                    except json.JSONDecodeError:
                        continue
                    delta = stream_delta_from_chunk(data)
                    if not delta:
                        continue
                    if not stream_started:
                        prefix_buffer += delta
                        stripped_prefix = prefix_buffer.lstrip()
                        if stripped_prefix.startswith("```") and "\n" not in stripped_prefix and len(stripped_prefix) < 24:
                            continue
                        delta = re.sub(r"^\s*```(?:markdown|md|text)?\s*", "", prefix_buffer, flags=re.I)
                        prefix_buffer = ""
                        stream_started = True
                    chunks.append(delta)
                    pending.append(delta)
                    await flush()
    except httpx.HTTPError as exc:
        add_event(job, "ai", f"Streaming indisponivel neste modelo/proxy; usando resposta completa ({exc}).", progress_start)
        content = await call_proxy(model, messages, max_tokens)
        add_report_delta(job, content, progress_end)
        return content
    await flush(force=True)
    return "".join(chunks)


def clean_model_text(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^```(?:markdown|md|text)?\s*", "", stripped)
        stripped = re.sub(r"\s*```$", "", stripped)
    return stripped.strip()


def extract_json_candidate(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^```(?:json)?\s*", "", stripped)
        stripped = re.sub(r"\s*```$", "", stripped)
    start = min([idx for idx in (stripped.find("{"), stripped.find("[")) if idx >= 0], default=-1)
    if start < 0:
        return stripped
    stack: list[str] = []
    in_string = False
    escape = False
    for index, char in enumerate(stripped[start:], start=start):
        if escape:
            escape = False
            continue
        if char == "\\":
            escape = True
            continue
        if char == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if char in "{[":
            stack.append("}" if char == "{" else "]")
        elif char in "}]":
            if not stack or stack[-1] != char:
                break
            stack.pop()
            if not stack:
                return stripped[start : index + 1]
    return stripped


def extract_string_field(text: str, field: str) -> str:
    match = re.search(rf'"{re.escape(field)}"\s*:\s*"((?:\\.|[^"\\])*)"', text, flags=re.S)
    if not match:
        return ""
    try:
        return json.loads(f'"{match.group(1)}"')
    except Exception:
        return match.group(1)


def parse_possible_json(text: str) -> Any:
    stripped = extract_json_candidate(text)
    try:
        return json.loads(stripped)
    except Exception as exc:
        summary = extract_string_field(stripped, "resumo_executivo") or extract_string_field(stripped, "resumo")
        fallback = {
            "raw": text,
            "parse_error": str(exc),
        }
        if summary:
            fallback["resumo_executivo"] = summary
        return fallback


LEGAL_SYSTEM_PROMPT = """Voce e um analista juridico brasileiro assistido por IA.
Trabalhe como consultor: analise profundamente, cite evidencia por pagina/bloco quando possivel e nao invente fatos.
Nao diga que substitui advogado. Seja util para triagem, revisao contratual e preparacao de decisao humana.
Responda em Markdown limpo, sem JSON, sem tabela grande e sem cercas de codigo."""


def deep_report_prompt(filename: str, text: str, extraction: dict[str, Any]) -> str:
    return f"""Gere um relatorio juridico profundo para o sistema MS RED PDF.

Arquivo: {filename}
Tipo: {extraction["document_type"].upper()}
Unidades extraidas: {extraction["total_pages"]} {extraction["unit_label"]}(s)
Caracteres extraidos: {extraction["total_chars"]}

O relatorio precisa cobrir TODO o escopo da proposta:
- leitura de contratos em PDF/DOCX;
- resumo automatico;
- sugestoes de clausulas;
- identificacao de riscos;
- historico da analise;
- pontos de compatibilidade com LGPD.

Estruture exatamente com estas secoes em Markdown:

# Analise juridica do documento
## 1. Conclusao executiva
Explique em linguagem de decisao: qual e a situacao do documento, nivel geral de risco e principais acoes recomendadas.

## 2. Resumo automatico
Resuma objeto, partes, prazo, valores, obrigacoes, multas, rescisao e pontos de dados pessoais. Se algo nao estiver no texto, diga que nao foi localizado.

## 3. Mapa de clausulas e obrigacoes
Liste as clausulas relevantes, citando pagina/bloco e trecho curto. Explique o efeito pratico de cada uma.

## 4. Riscos identificados
Classifique riscos em Alto, Medio e Baixo. Para cada risco, informe evidencia, impacto, causa e recomendacao.

## 5. Sugestoes de clausulas e melhorias
Proponha redacoes alternativas ou complementares. Inclua texto sugerido quando fizer sentido.

## 6. LGPD e privacidade
Avalie dados pessoais, base legal, finalidade, compartilhamento, seguranca, retencao, direitos do titular e lacunas de conformidade.

## 7. Perguntas para o cliente
Liste perguntas objetivas que faltam para fechar a analise.

## 8. Checklist de proximos passos
Liste acoes praticas em ordem de prioridade.

## 9. Historico da analise
Explique como a analise foi feita: formato recebido, extracao, estrategia, modelo usado e limites observados.

Regras:
- Nao crie fatos nao presentes no documento.
- Se a evidencia estiver fraca, escreva \"nao localizado no documento\".
- Use bullets curtos, mas com profundidade.
- Evite JSON.
- Cite pagina/bloco no formato [Pagina 1] ou [Bloco 2].

Documento:
{text}"""


def page_prompt(filename: str, page_number: int, text: str) -> str:
    return f"""Analise somente esta unidade do documento juridico e produza notas tecnicas em Markdown.

Arquivo: {filename}
Unidade: {page_number}

Cubra:
- resumo da unidade;
- clausulas/obrigacoes relevantes;
- riscos com gravidade;
- sugestoes de melhoria;
- pontos de LGPD;
- evidencias curtas.

Texto:
{text}"""


def consolidation_prompt(filename: str, extraction: dict[str, Any], page_notes: list[Any]) -> str:
    notes = json.dumps(page_notes, ensure_ascii=False)[:MAX_CHARS_WHOLE]
    return f"""Consolide as notas por unidade em um relatorio juridico unico, profundo e legivel.

Arquivo: {filename}
Tipo: {extraction["document_type"].upper()}
Unidades extraidas: {extraction["total_pages"]} {extraction["unit_label"]}(s)
Caracteres extraidos: {extraction["total_chars"]}

O relatorio precisa cobrir TODO o escopo da proposta:
- leitura de contratos em PDF/DOCX;
- resumo automatico;
- sugestoes de clausulas;
- identificacao de riscos;
- historico da analise;
- compatibilidade com LGPD.

Use exatamente estas secoes:

# Analise juridica do documento
## 1. Conclusao executiva
## 2. Resumo automatico
## 3. Mapa de clausulas e obrigacoes
## 4. Riscos identificados
## 5. Sugestoes de clausulas e melhorias
## 6. LGPD e privacidade
## 7. Perguntas para o cliente
## 8. Checklist de proximos passos
## 9. Historico da analise

Regras:
- Markdown limpo, sem JSON.
- Nao invente fatos.
- Diferencie risco alto, medio e baixo.
- Inclua texto sugerido de clausulas quando fizer sentido.
- Cite pagina/bloco sempre que houver evidencia.

Notas por unidade:
{notes}"""


async def analyze_whole(job: Job, extraction: dict[str, Any]) -> dict[str, Any]:
    marker = "Pagina" if extraction["unit_label"] == "pagina" else "Bloco"
    all_text = "\n\n".join(f"[{marker} {p['page']}]\n{p['text']}" for p in extraction["pages"] if p["text"])
    all_text = trim_text(redact_for_lgpd(all_text), MAX_CHARS_WHOLE)
    job.report_draft = ""
    add_event(job, "ai", f"Gerando relatorio juridico profundo com {job.model}.", 46, {"chars": len(all_text)})
    add_event(job, "report-reset", "Streaming do relatorio iniciado.", 47)
    content = await call_proxy_stream(
        job.model,
        [
            {"role": "system", "content": LEGAL_SYSTEM_PROMPT},
            {"role": "user", "content": deep_report_prompt(job.filename, all_text, extraction)},
        ],
        job=job,
        max_tokens=REPORT_TOKENS,
        progress_start=48,
        progress_end=82,
    )
    report = clean_model_text(content)
    job.report_draft = report
    add_event(job, "ai", "Relatorio consultivo recebido.", 82)
    return {"strategy": "whole", "report_markdown": report}


async def analyze_pages(job: Job, extraction: dict[str, Any]) -> dict[str, Any]:
    pages = [page for page in extraction["pages"] if page["text"]]
    capped = pages[:MAX_ANALYSIS_PAGES]
    if len(pages) > len(capped):
        add_event(job, "strategy", f"Analise limitada as primeiras {MAX_ANALYSIS_PAGES} paginas por configuracao.", 38)
    page_notes: list[Any] = []
    total = max(len(capped), 1)
    job.report_draft = ""
    add_event(job, "report-reset", "Streaming das notas preliminares por unidade iniciado.", 39)
    for idx, page in enumerate(capped, start=1):
        text = trim_text(redact_for_lgpd(page["text"]), MAX_CHARS_PER_PAGE)
        progress = 40 + int((idx - 1) / total * 38)
        add_event(job, "ai-page", f"Analisando pagina {page['page']} ({idx}/{total}) com {job.model}.", progress)
        try:
            add_report_delta(job, f"\n\n## Notas preliminares da unidade {page['page']}\n", progress)
            content = await call_proxy_stream(
                job.model,
                [
                    {"role": "system", "content": LEGAL_SYSTEM_PROMPT},
                    {"role": "user", "content": page_prompt(job.filename, page["page"], text)},
                ],
                job=job,
                max_tokens=PAGE_NOTE_TOKENS,
                progress_start=progress,
                progress_end=progress + 1,
            )
            page_notes.append({"unidade": page["page"], "notas": clean_model_text(content)})
            add_event(job, "ai-page", f"Pagina {page['page']} analisada.", progress + 1)
        except Exception as exc:
            page_notes.append({"unidade": page["page"], "erro": str(exc)})
            add_event(job, "ai-page", f"Falha na unidade {page['page']}: {exc}", progress + 1)
    add_event(job, "ai", f"Consolidando {len(page_notes)} analise(s) por pagina.", 82)
    job.report_draft = ""
    add_event(job, "report-reset", "Streaming do relatorio consolidado iniciado.", 83)
    consolidation_messages = [
        {"role": "system", "content": LEGAL_SYSTEM_PROMPT},
        {"role": "user", "content": consolidation_prompt(job.filename, extraction, page_notes)},
    ]
    review_model = REVIEW_MODEL or job.model
    try:
        content = await call_proxy_stream(
            review_model,
            consolidation_messages,
            job=job,
            max_tokens=REPORT_TOKENS,
            progress_start=84,
            progress_end=94,
        )
    except httpx.HTTPError as exc:
        if review_model == job.model:
            raise
        add_event(job, "ai", f"Modelo de revisao indisponivel ({exc}); consolidando com {job.model}.", 84)
        job.report_draft = ""
        add_event(job, "report-reset", f"Streaming do relatorio consolidado iniciado com {job.model}.", 85)
        content = await call_proxy_stream(
            job.model,
            consolidation_messages,
            job=job,
            max_tokens=REPORT_TOKENS,
            progress_start=86,
            progress_end=94,
        )
    report = clean_model_text(content)
    job.report_draft = report
    return {
        "strategy": "pages",
        "pages_analyzed": len(page_notes),
        "page_notes": page_notes,
        "report_markdown": report,
    }


async def process_job(job_id: str) -> None:
    job = jobs[job_id]
    try:
        job.status = "running"
        add_event(job, "start", "Analise iniciada.", 5)
        path = Path(job.file_path or "")
        extraction = await asyncio.to_thread(extract_document, path, job)
        if pdf_needs_ocr(extraction):
            extraction = await asyncio.to_thread(apply_pdf_ocr, path, job, extraction)
        if extraction["total_chars"] == 0:
            if extraction["document_type"] == "pdf" and not ENABLE_OCR:
                raise RuntimeError("Nenhum texto foi extraido. O PDF parece escaneado e o OCR esta desativado.")
            raise RuntimeError("Nenhum texto foi extraido mesmo apos OCR. O documento pode estar ilegivel, protegido ou com imagem muito ruim.")
        if extraction["empty_pages"]:
            add_event(job, "extract", f"{extraction['empty_pages']} unidade(s) ainda ficaram com pouco texto apos extracao/OCR.", 35)
        strategy = choose_strategy(job.mode, extraction)
        add_event(job, "strategy", f"Estrategia definida: {'documento inteiro' if strategy == 'whole' else 'pagina por pagina'}.", 38, {"strategy": strategy})
        if strategy == "whole":
            analysis = await analyze_whole(job, extraction)
        else:
            analysis = await analyze_pages(job, extraction)
        result = {
            "job_id": job.id,
            "filename": job.filename,
            "model": job.model,
            "review_model": REVIEW_MODEL,
            "mode": job.mode,
            "created_at": job.created_at,
            "finished_at": time.time(),
            "redacted_before_ai": REDACT_BEFORE_AI,
            "document": {
                "type": extraction["document_type"],
                "unit_label": extraction["unit_label"],
                "pages": extraction["total_pages"],
                "chars": extraction["total_chars"],
                "metadata": extraction["metadata"],
                "empty_pages": extraction["empty_pages"],
                "ocr_applied": extraction.get("ocr_applied", False),
                "ocr_pages": extraction.get("ocr_pages", []),
                "ocr_lang": extraction.get("ocr_lang", ""),
                "ocr_dpi": extraction.get("ocr_dpi", ""),
            },
            **analysis,
        }
        result_path = RESULT_DIR / f"{job.id}.json"
        result_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        job.result = result
        job.status = "completed"
        add_event(job, "history", "Historico da analise salvo.", 94)
        add_event(job, "done", "Relatorio final pronto.", 100)
    except Exception as exc:
        job.status = "failed"
        job.error = str(exc)
        add_event(job, "error", f"Analise interrompida: {exc}", 100)


@app.on_event("startup")
async def startup() -> None:
    ensure_dirs()


@app.get("/")
async def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/healthz")
async def healthz() -> dict[str, Any]:
    return {"ok": True, "service": "msredpdf", "model": DEFAULT_MODEL, "proxy": PROXY_BASE_URL}


@app.get("/api/models")
async def models() -> dict[str, Any]:
    return {"default": DEFAULT_MODEL, "review": REVIEW_MODEL, "recommended": RECOMMENDED_MODELS}


@app.get("/api/history")
async def history() -> dict[str, Any]:
    ensure_dirs()
    items: list[dict[str, Any]] = []
    for path in sorted(RESULT_DIR.glob("*.json"), key=lambda item: item.stat().st_mtime, reverse=True)[:40]:
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            continue
        document = data.get("document", {})
        items.append(
            {
                "id": data.get("job_id", path.stem),
                "filename": data.get("filename", "documento"),
                "model": data.get("model", ""),
                "strategy": data.get("strategy", ""),
                "document_type": document.get("type", ""),
                "units": document.get("pages", 0),
                "created_at": data.get("created_at"),
                "finished_at": data.get("finished_at"),
            }
        )
    return {"items": items}


@app.post("/api/jobs")
async def create_job(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    model: str = Form(DEFAULT_MODEL),
    mode: str = Form("auto"),
) -> JSONResponse:
    ensure_dirs()
    if mode not in {"auto", "whole", "pages"}:
        raise HTTPException(status_code=400, detail="modo invalido")
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="envie um PDF ou DOCX")
    job_id = uuid.uuid4().hex[:16]
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(file.filename).name)[:160] or f"documento{suffix}"
    target = UPLOAD_DIR / f"{job_id}-{safe_name}"
    size = 0
    max_bytes = MAX_UPLOAD_MB * 1024 * 1024
    with target.open("wb") as out:
        while chunk := await file.read(1024 * 1024):
            size += len(chunk)
            if size > max_bytes:
                out.close()
                target.unlink(missing_ok=True)
                raise HTTPException(status_code=413, detail=f"PDF acima de {MAX_UPLOAD_MB} MB")
            out.write(chunk)
    job = Job(id=job_id, filename=safe_name, model=model or DEFAULT_MODEL, mode=mode, document_type=suffix.lstrip("."), file_path=str(target))
    jobs[job_id] = job
    add_event(job, "upload", f"Documento recebido: {safe_name} ({size / 1024 / 1024:.2f} MB).", 3)
    background_tasks.add_task(process_job, job_id)
    return JSONResponse(public_job(job), status_code=202)


@app.get("/api/jobs/{job_id}")
async def get_job(job_id: str) -> dict[str, Any]:
    job = jobs.get(job_id)
    if not job:
        result_path = RESULT_DIR / f"{job_id}.json"
        if result_path.exists():
            try:
                return public_saved_job(job_id, json.loads(result_path.read_text(encoding="utf-8")))
            except Exception:
                pass
        raise HTTPException(status_code=404, detail="job nao encontrado")
    return public_job(job)


@app.get("/api/jobs/{job_id}/events")
async def job_events(job_id: str, request: Request) -> StreamingResponse:
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="job nao encontrado")

    async def stream() -> AsyncIterator[str]:
        cursor = 0
        while True:
            job = jobs.get(job_id)
            if not job:
                break
            while cursor < len(job.events):
                event = job.events[cursor]
                cursor += 1
                yield f"id: {event['seq']}\nevent: progress\ndata: {json.dumps(event, ensure_ascii=False)}\n\n"
            if job.status in {"completed", "failed"}:
                yield f"event: final\ndata: {json.dumps(public_job(job), ensure_ascii=False)}\n\n"
                break
            if await request.is_disconnected():
                break
            await asyncio.sleep(0.35)

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )
